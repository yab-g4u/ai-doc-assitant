import streamlit as st
import os
from uuid import uuid4

from src.loader import load_document
from src.splitter import split_documents
from src.vectorstore import store_embeddings
from src.memory import create_memory
from src.gemini_llm import generate_gemini_response

from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
import logging
logging.basicConfig(level=logging.INFO)
logging.info("App starting..") 

from PyPDF2 import PdfReader
import docx

st.set_page_config(page_title="AI Doc Assistant", layout="wide", initial_sidebar_state="expanded")
GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]

# --- Directories ---
UPLOAD_DIR = "data/uploads"
INDEX_DIR = "data/indexes"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(INDEX_DIR, exist_ok=True)

# --- File Readers ---
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def read_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

# --- SIDEBAR ---
with st.sidebar:
    st.title("📂 Upload & Settings")
    uploaded_file = st.file_uploader(
        "Upload your document (.pdf, .txt, .docx)",
        type=["pdf", "txt", "docx"]
    )
    st.markdown("---")
    st.subheader("ℹ️ About")
    st.info(
        "AI Doc Assistant lets you chat with your documents using Gemini LLM.\n\n"
        "Upload a PDF, TXT, or DOCX file and ask questions about its content."
    )

# --- MAIN CHAT INTERFACE ---
st.markdown(
    """
    <style>
    .chat-container {height: 65vh; overflow-y: auto; padding: 1rem; background: #f7f7fa; border-radius: 10px;}
    .user-msg {background: #d1e7dd; color: #222; padding: 0.7em 1em; border-radius: 12px 12px 0 12px; margin-bottom: 0.5em; max-width: 80%; align-self: flex-end;}
    .bot-msg {background: #fff; color: #222; padding: 0.7em 1em; border-radius: 12px 12px 12px 0; margin-bottom: 0.5em; max-width: 80%; align-self: flex-start; border: 1px solid #eee;}
    .chat-history {display: flex; flex-direction: column;}
    .typing {font-style: italic; color: #888;}
    </style>
    """,
    unsafe_allow_html=True,
)

# --- Session State Setup ---
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid4())
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "retriever" not in st.session_state:
    st.session_state.retriever = None

# --- Handle file upload and processing ---
if uploaded_file:
    file_path = os.path.join(UPLOAD_DIR, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    st.sidebar.success("✅ File uploaded successfully.")

    with st.spinner("Processing document..."):
        ext = uploaded_file.name.lower().split(".")[-1]
        if ext == "pdf":
            content = read_pdf(file_path)
        elif ext == "txt":
            content = read_txt(file_path)
        elif ext == "docx":
            content = read_docx(file_path)
        else:
            content = None

        if not content or not content.strip():
            st.error("❌ The document could not be processed or is empty. Please try another file.")
        else:
            docs = [Document(page_content=content)]
            chunks = split_documents(docs)
            index_path = os.path.join(INDEX_DIR, st.session_state.session_id)
            vectorstore = store_embeddings(chunks, index_path)
            retriever = vectorstore.as_retriever(search_kwargs={"k": 4})
            st.session_state.vectorstore = vectorstore
            st.session_state.retriever = retriever

# --- CHAT UI ---
st.title("🧠 Chat With Your Document")

chat_placeholder = st.empty()

def render_chat():
    messages = st.session_state.chat_history
    chat_html = '<div class="chat-container"><div class="chat-history">'
    for q, a in messages:
        chat_html += f'<div class="user-msg">🧑 {q}</div>'
        chat_html += f'<div class="bot-msg">🤖 {a}</div>'
    chat_html += '</div></div>'
    chat_placeholder.markdown(chat_html, unsafe_allow_html=True)

render_chat()

# --- INPUT BOX AT BOTTOM ---
with st.form(key="chat_form", clear_on_submit=True):
    user_input = st.text_input("Your question:", placeholder="Type your question and press Enter...", label_visibility="collapsed")
    submitted = st.form_submit_button("Send", use_container_width=True)

# --- HANDLE USER INPUT ---
if submitted and user_input and st.session_state.retriever:
    with st.spinner("🤖 Thinking..."):
        def answer_query(query):
            docs = st.session_state.retriever.get_relevant_documents(query)
            context = "\n".join([doc.page_content for doc in docs])
            prompt = f"Use the following context to answer the question:\n\n{context}\n\nQuestion: {query}"
            return generate_gemini_response(prompt)
        answer = answer_query(user_input)
        st.session_state.chat_history.append((user_input, answer))
        render_chat()


