from pathlib import Path
from langchain.schema import Document

def load_document(file_path):
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
        return [Document(page_content=text, metadata={"source": file_path})]
    
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": file_path})]
    
    elif ext == ".docx":
        import docx
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return [Document(page_content=text, metadata={"source": file_path})]
    
    else:
        raise ValueError("Unsupported file type. Only .pdf, .txt, and .docx are allowed.")
